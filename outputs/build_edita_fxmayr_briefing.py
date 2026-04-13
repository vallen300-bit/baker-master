"""Build Edita briefing DOCX — What DV communicated to AO at FX Mayr, 11 April 2026."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# --- Header ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BRIEFING FOR EDITA')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('What DV Communicated to AO at FX Mayr — 11 April 2026')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Source: Fireflies audio recording, translated from Russian\nPrepared by Baker AO PM — 12 April 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()  # spacer

# --- Purpose ---
p = doc.add_heading('Purpose', level=1)
doc.add_paragraph(
    'DV walked AO through the financing table line by line at FX Mayr on April 11 (24-minute conversation in Russian). '
    'This document captures exactly what AO was told, organized by topic, including his questions and DV\'s answers. '
    'The formal capital call letter and Excel must correspond to both this briefing and v009.'
)
p = doc.add_paragraph()
run = p.add_run('Rule: ')
run.bold = True
p.add_run('If the capital call letter contradicts v009 → Edita\'s table wins. '
          'The letter must be consistent with what AO heard, but numbers must match v009.')

doc.add_paragraph()

# ============================================================
# ISSUE 1
# ============================================================
doc.add_heading('Issue 1: Construction Costs (Rows 14–16)', level=1)

doc.add_heading('What AO Was Told', level=2)
items = [
    'Hagenauer contract: EUR 121.5M (floating, can decrease) + EUR 17M fees (separate supervision contract).',
    'Paid to Hagenauer: EUR 114M + all fees.',
    'May 2025 gap: Hagenauer self-financed construction for 2–3 months when DV couldn\'t pay ~EUR 9M. '
    'Then Hagenauer took ~EUR 5M from his "general fund" as compensation.',
    'Delay cost: EUR 1,341K additional (Oct–Dec actual vs Sep projection) — workers, extra two months.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Replacement Costs (Rows 15–16) — Explained in Detail', level=2)
items = [
    'Electrical audit: independent consultant EUR 250–300K to check all certificates before May 26 Vienna inspection.',
    'Kitchen electrical shock incident disclosed to AO.',
    'MEP: no plans available (administrator refuses to hand over). Ventilation/temperature issues. Sewage smell in sold apartments.',
    'Must fix before May 26 or risk permit revocation + press crisis.',
    'These are costs to redo work at own expense because Hagenauer is bankrupt — no warranty claims possible.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('AO\'s Reaction', level=2)
doc.add_paragraph('"I understand. All clear." — No pushback on these items.')

doc.add_paragraph()

# ============================================================
# ISSUE 2
# ============================================================
doc.add_heading('Issue 2: VAT / NDS Clawback (Row 28)', level=1)

doc.add_heading('What AO Was Told', level=2)
items = [
    'November 20, 2025: Austria adopted Betrugsbekämpfungsgesetz — new law targeting luxury apartments.',
    'Apartments above EUR 2.5M no longer qualify for business-use NDS exemption.',
    'Triggered by Benko bankruptcies — he sold properties as "business" to avoid NDS.',
    'Impact on us: EUR 5,770K NDS obligation.',
    'Previously paid EUR 373K NDS (for the "private use" half under old rules).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('AO\'s Question', level=2)
p = doc.add_paragraph()
run = p.add_run('AO: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x99)
p.add_run('"The law does not have retroactive power?"')

doc.add_heading('DV\'s Answer', level=2)
doc.add_paragraph(
    '"We should not pay them either way. We wrote a paper that money comes to us from sales and we do not pay NDS. '
    'But when they introduced the law, they came to us." '
    'DV explained: took EUR 4,430K from apartment sale proceeds to pay NDS. '
    'Bank doesn\'t know yet. Lawyers say: (a) DV had legal obligation to pay NDS, '
    '(b) bank contract didn\'t include NDS clawback provision. "The bank will grumble and scream, but we have legal basis."'
)

doc.add_heading('AO\'s Reaction', level=2)
doc.add_paragraph('"How much did they get?" — factual question, no objection. Accepted the explanation.')

doc.add_paragraph()

# ============================================================
# ISSUE 3
# ============================================================
doc.add_heading('Issue 3: Fund\'s EUR 5.7M — Cashed Bank Guarantee', level=1)

p = doc.add_paragraph()
run = p.add_run('⚠ CRITICAL — This changes how the capital call should be framed.')
run.bold = True
run.font.color.rgb = RGBColor(0xcc, 0x00, 0x00)

doc.add_heading('What AO Was Told', level=2)
items = [
    'The Fund provided a bank guarantee (~EUR 5.5–5.6M) at the start of construction.',
    'This was NOT a capital call — it was a guarantee instrument.',
    'At some point DV cashed it ("broke the guarantee") to cover delay expenses.',
    'Fund "didn\'t feel anything when they cashed in — because I am in charge of everything."',
    'The guarantee became cash and paid for: Mandarin delay, pre-opening costs, works.',
    'There was also a prior imbalance of EUR ~3.3M that AO had previously repaid. DV has agreement to return EUR 3M to the Fund from this.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('AO\'s Questions', level=2)
p = doc.add_paragraph()
run = p.add_run('AO: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x99)
p.add_run('"It turns out it was not you who gave more, but the fund gave more?"')

p = doc.add_paragraph()
run = p.add_run('DV: ')
run.bold = True
p.add_run('"Yes. There is one more nuance..." [explained the guarantee conversion and the EUR 3.3M prior imbalance]')

p = doc.add_paragraph()
run = p.add_run('AO: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x99)
p.add_run('"Did they not feel anything when they cashed in?"')

p = doc.add_paragraph()
run = p.add_run('DV: ')
run.bold = True
p.add_run('"No. Because I am in charge of everything."')

doc.add_heading('Implication for Edita', level=2)
doc.add_paragraph(
    'Row 35 in v009 shows "Fund (52%) Actual: 5,738,339." AO understands this as a cashed guarantee, '
    'not an equity contribution. The capital call letter should NOT frame this as '
    '"the Fund over-invested and you must rectify." '
    'Correct framing: "The guarantee was converted to cover the delay. Going forward, both sides contribute."'
)

doc.add_paragraph()

# ============================================================
# ISSUE 4
# ============================================================
doc.add_heading('Issue 4: Total Deficit & Capital Call', level=1)

doc.add_heading('What AO Was Told', level=2)

# Table
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

cells = table.rows[0].cells
cells[0].text = 'Item'
cells[1].text = 'Amount'
for cell in cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

data = [
    ('Sep 2025 capital call (covered)', 'EUR 44,221,154'),
    ('Oct–Dec deficit (delay)', 'EUR 5,848K'),
    ('Jan–Jun 2026 projected', 'EUR 21M'),
    ('Total uses Oct–Jun', 'EUR 27,134K'),
]
for i, (item, amount) in enumerate(data):
    cells = table.rows[i + 1].cells
    cells[0].text = item
    cells[1].text = amount

doc.add_paragraph()
doc.add_paragraph('DV told AO:')
items = [
    '"Total forward deficit: approximately 11 million."',
    '"It is necessary to divide it in half. That is, it is six. Until June."',
    '"They do not need everything right now. They will need two, two and two."',
    '"From Monday I start working with all banks. We need a pillow. Not to take from shareholders."',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.runs[0].italic = True

doc.add_heading('Director\'s Decision (12 April)', level=2)
p = doc.add_paragraph()
run = p.add_run('Final capital call to AO = EUR 7M')
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph()

table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

cells = table.rows[0].cells
cells[0].text = 'Month'
cells[1].text = 'AO Transfers'
for cell in cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

data = [('April', 'EUR 2,500,000'), ('May', 'EUR 2,500,000'), ('June', 'EUR 2,000,000')]
for i, (month, amount) in enumerate(data):
    cells = table.rows[i + 1].cells
    cells[0].text = month
    cells[1].text = amount

doc.add_paragraph()
doc.add_paragraph(
    'AO heard "~6M" at FX Mayr. The letter should frame the EUR 1M difference as contingency '
    'for items still being finalized (MEP audit results, subcontractor negotiations, bank interest true-up).'
)

doc.add_paragraph()

# ============================================================
# ISSUE 5
# ============================================================
doc.add_heading('Issue 5: EUR 2.5M Special Deal (Outside Framework)', level=1)

doc.add_heading('What AO Was Told', level=2)
p = doc.add_paragraph()
run = p.add_run('DV: ')
run.bold = True
p.add_run('"Two and a half which you transferred. I explain to them [the Fund]: '
          'I finance it. Two and a half. And I solve the problems myself. Close it by four. '
          'That is, we will take into account as four."')

doc.add_paragraph()
items = [
    'AO\'s EUR 2.5M sits OUTSIDE the 48/52 framework.',
    'AO gets a EUR 4M return claim at distribution (not EUR 2.5M shareholder capital).',
    'The Fund was told: "we closed the construction gap for 4M, carry on."',
    '"It was not taken into account in our relations" — meaning Fund/AO relations.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('AO\'s Question', level=2)
p = doc.add_paragraph()
run = p.add_run('AO: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0x99)
p.add_run('"These 2,500 — you want to show that he [has them] at you...?"')

p = doc.add_paragraph()
run = p.add_run('DV: ')
run.bold = True
p.add_run('Explained the 2.5M → 4M return structure. AO accepted without objection.')

doc.add_heading('Implication for Edita', level=2)
doc.add_paragraph(
    'v009 Row 36 shows EUR 2.5M as "AO (48%) Actual." But DV told AO this is outside the framework. '
    'Question: Should Row 36 show EUR 0 for AO (with the 2.5M on a separate line), '
    'or keep as-is with a footnote? The capital call letter must be consistent with AO\'s understanding.'
)

doc.add_paragraph()

# ============================================================
# ISSUE 6
# ============================================================
doc.add_heading('Issue 6: Hagenauer Insolvency / EUR 19M Claim', level=1)

doc.add_heading('What AO Was Told', level=2)
items = [
    'Insolvency administrator has EUR 19M in claims from subcontractors.',
    'Only asset of Hagenauer = claim against RG7 for EUR 19M.',
    'Risk: a financial buyer could purchase claims for ~EUR 3M and litigate for years.',
    'Option: pay EUR 2M quick settlement and close.',
    'Warranty strategy: approach major subcontractors (MEP), offer EUR 500K + defect repair for 3-year guarantee with bank backing.',
    'EUR 4M warranty/settlement fund covers this strategy.',
    '"It is cheaper to negotiate with them and get a 3-year guarantee than to litigate for 10 years and sell the hotel at a discount."',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('AO\'s Reaction', level=2)
doc.add_paragraph('"Okay. All clear." — Accepted the logic.')

doc.add_paragraph()

# ============================================================
# ACTION ITEMS
# ============================================================
doc.add_heading('Action Items for Edita', level=1)

actions = [
    ('v009 Alignment', 'Does the table match what DV told AO? Key check: EUR 2.5M in Row 36 — inside or outside the 48/52 framework?'),
    ('Fund Guarantee', 'Row 35 shows Fund EUR 5,738K as "Actual." AO understands this as a cashed guarantee. How should the capital call letter present this?'),
    ('Capital Call Column R', 'Adjust to reflect EUR 7M to AO: EUR 2.5M April, EUR 2.5M May, EUR 2M June.'),
    ('Fund Additional EUR 2,383K', 'DV told AO he\'d figure this out by Sunday. Has it been resolved? What is the source?'),
    ('Row 20 (Working Capital)', 'DV mentioned during your working session that Q20 should be either EUR 1.5M or EUR 600K. Please confirm the correct figure.'),
]

table = doc.add_table(rows=len(actions) + 1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

cells = table.rows[0].cells
cells[0].text = '#'
cells[1].text = 'Item'
cells[2].text = 'Question / Action'
for cell in cells:
    for p in cell.paragraphs:
        p.runs[0].bold = True

for i, (item, question) in enumerate(actions):
    cells = table.rows[i + 1].cells
    cells[0].text = str(i + 1)
    cells[1].text = item
    cells[2].text = question

doc.add_paragraph()

# --- Footer ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Briefing —')
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.size = Pt(9)

# Save
outpath = '/Users/dimitry/Desktop/EDITA_BRIEFING_AO_FXMAYR_11APR2026.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
